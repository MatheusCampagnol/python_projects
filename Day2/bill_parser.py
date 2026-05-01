#Testing Parsing Methods. 
#Importing dependencies.
import argparse
from docx import Document


def get_user_input():
    try:
        total_bill = float(input("What was the total bill? "))
        if total_bill <= 0:
            raise ValueError("Total value of bill must be greater than zero.")

        tip_percentage = float(input("How much tip would you like to give? 10, 12, or 15? "))
        if tip_percentage < 0:
            raise ValueError("Tip percentage cannot be negative.")

        number_of_people = int(input("How many people to split the bill? "))
        if number_of_people <= 0:
            raise ValueError("Number of people must be greater than zero.")

        return total_bill, tip_percentage, number_of_people

    except ValueError as e:
        print(f"Error: {e}")
        exit(1)


def bill_calculation(total_bill, tip_percentage, number_of_people):
    try:
        value_per_person = (total_bill + (total_bill * tip_percentage / 100)) / number_of_people
        print(f"Each person should pay: ${value_per_person:.2f}")
        return value_per_person
    except ZeroDivisionError as e:
        print(f"Error: {e}")
        exit(1)


def document_processing(total_bill, tip_percentage, number_of_people, value_per_person):
    doc = Document()
    doc.add_paragraph("Hello! This is a bill splitter program.")
    doc.add_paragraph(f"Total Bill: ${total_bill:.2f}")
    doc.add_paragraph(f"Tip Percentage: {tip_percentage}%")
    doc.add_paragraph(f"Number of People: {number_of_people}")
    doc.add_paragraph(f"Each person should pay: ${value_per_person:.2f}")
    doc.save("bill_splitter_output.docx")
    print("Document 'bill_splitter_output.docx' has been created with the bill details.")


def main():
    parser = argparse.ArgumentParser(description="Split a bill between people with a tip percentage.")
    parser.add_argument("total_bill", nargs="?", type=float, help="Total bill amount")
    parser.add_argument("tip_percentage", nargs="?", type=float, help="Tip percentage to add")
    parser.add_argument("number_of_people", nargs="?", type=int, help="Number of people to split the bill")
    args = parser.parse_args()

    if args.total_bill is not None and args.tip_percentage is not None and args.number_of_people is not None:
        total_bill = args.total_bill
        tip_percentage = args.tip_percentage
        number_of_people = args.number_of_people
    else:
        total_bill, tip_percentage, number_of_people = get_user_input()

    value_per_person = bill_calculation(total_bill, tip_percentage, number_of_people)
    document_processing(total_bill, tip_percentage, number_of_people, value_per_person)


if __name__ == "__main__":
    main()
