#Importing dependencies.
from docx import Document


#Variables:
total_bill = 0.0
tip_percentage = 0.0
number_of_people = 0
value_per_person = 0.0

#Function to get user input.

def get_user_input():
    
    global total_bill, tip_percentage, number_of_people
  
#Tries to get input, and in case of error throws a ValueError with a message and exits the program.  
    try:  
        total_bill = float(input("What was the total bill?"))
        if total_bill <= 0:
            raise ValueError("Total value of bill must be greater than zero.")
        
        tip_percentage = float(input("How much tip would you like to give? 10, 12, or 15?"))
        if tip_percentage < 0:
            raise ValueError("Tip percentage cannot be negative.")
        
        number_of_people = int(input("How many people to split the bill?"))
        if number_of_people <= 0:
            raise ValueError("Number of people must be greater than zero.")
        
#ValueError is assigned to e and printed.        
    except ValueError as e:
        print(f"Error: {e}")
        exit(1)


def bill_calculation(total_bill, tip_percentage, number_of_people):

    try:
        value_per_person = (total_bill + (total_bill * tip_percentage / 100)) / number_of_people
        print(f"Each person should pay: ${value_per_person:.2f}")
    except ZeroDivisionError as e:
        print(f"Error: {e}")

def document_processing():
#Creates the new document.
    doc = Document()
    doc.add_paragraph("Hello! This is a bill splitter program.")
    doc.add_paragraph(f"Total Bill: ${total_bill:.2f}")
    doc.add_paragraph(f"Tip Percentage: {tip_percentage}%")
    doc.add_paragraph(f"Number of People: {number_of_people}")
    doc.add_paragraph(f"Each person should pay: ${value_per_person:.2f}")
    doc.save("bill_splitter_output.docx")
#Confirmation message that the document has been created.
    print("Document 'bill_splitter_output.docx' has been created with the bill details.")   

def main():

#Get user input.
    get_user_input()

#Calculate each person's share of the bill.    
    bill_calculation(total_bill, tip_percentage, number_of_people)
#Creates the document for the guests.
    document_processing()
    
#Calls Main Funcition. Start of the program.
if __name__ == "__main__":
    main()
    